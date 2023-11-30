from abc import ABC, abstractmethod
import json
import fitz
import pandas as pd

from docx import Document as DocxDocumentOBJ

from document_processor import DocumentProcessor


class Document(ABC):
    def __init__(self, path, document):
        self.document = document
        self.path = path

    @classmethod
    def fromPath(cls, filepath):
        try:
            document = cls.open(
                filepath
            )  # Call the open method which is defined in the subclass
            return cls(filepath, document)  # Return an instance of the subclass
        except Exception as e:
            print(f"Error opening file {filepath}: {e}")
            return None

    @staticmethod
    @abstractmethod
    def open(filepath):
        pass

    @abstractmethod
    def process(self):
        pass

    def save_to_parquet(self, filepath, data_json=None, log=True):
        # try to cast data_json as list if unsuccessful raise

        # Convert to DataFrame
        df = pd.DataFrame(data_json, columns=["chunks"])

        # Write to Parquet

        df.to_parquet(filepath)

        return filepath

    def save(self, filepath, data_json=None, log=True):
        if log:
            print(f"Saving to {filepath}")

        if data_json is None:
            raise ValueError("Empty content cannot be saved")

        self.save_to_parquet(filepath, data_json, log)

        # with open(filepath, "w", encoding="utf-8") as file:
        #    json.dump(data_json, file, indent=4, ensure_ascii=False)

    def process_and_save(self, filepath, log=True):
        try:
            data_json = self.process()
            self.save(filepath, data_json, log)

            return True
        except Exception as e:
            if log:
                print(f"Error processing file {filepath}: {e}")
            return False


class PDFDocument(Document):
    def open(filepath):
        # Use fitz to open the PDF file
        return fitz.open(filepath)

    def process(self):
        document_processor = DocumentProcessor()

        all_data = document_processor.extract_content(
            self.document
        )  # Assuming extract_content is defined elsewhere

        if not all_data:
            raise ValueError("Empty PDF Content")

        extracted_data = document_processor.postprocess(
            all_data, "pdf"
        )  # Assuming postprocess is defined elsewhere
        return extracted_data


class EPUBDocument(Document):
    def open(filepath):
        # Use fitz to open the EPUB file
        return fitz.open(filepath)

    def process(self):
        document_processor = DocumentProcessor()

        all_data = document_processor.extract_content(
            self.document
        )  # Assuming extract_content is defined elsewhere

        if not all_data:
            raise ValueError("Empty EPUB/MOBI Content")

        extracted_data = document_processor.postprocess(
            all_data, "epub"
        )  # Assuming postprocess is defined elsewhere
        return extracted_data


class DOCXDocument(Document):
    def open(filepath):
        # Open the DOCX file using a DOCX document object
        return Document(
            filepath
        )  # Assuming Document here refers to a library for handling DOCX files

    def process(self):
        document_processor = DocumentProcessor()

        extracted_data = document_processor.postprocess(
            self.document.paragraphs, "docx"
        )  # Assuming postprocess is defined elsewhere

        if not extracted_data:
            raise ValueError("Empty DOCX Content")

        return extracted_data


class TXTDocument(Document):
    def open(filepath):
        # Open the text file
        with open(filepath, "r", encoding="utf-8") as file:
            return file.read()

    def process(self):
        document_processor = DocumentProcessor()

        extracted_data = document_processor.postprocess(
            self.document, "txt"
        )  # Assuming postprocess is defined elsewhere
        if not extracted_data:
            raise ValueError("Empty TXT Content")
        return extracted_data


class DJVUDocument(Document):
    def open(filepath):
        document_processor = DocumentProcessor()
        # Open the DJVU file using a method to read DJVU files
        return document_processor.read_DJVU(
            filepath
        )  # Assuming read_DJVU is defined elsewhere

    def process(self):
        document_processor = DocumentProcessor()

        # Process the DJVU document
        extracted_data = document_processor.postprocess(
            self.document, "djvu"
        )  # Assuming postprocess is defined elsewhere
        if not extracted_data:
            raise ValueError("Empty DJVU Content")
        return extracted_data
